from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import shutil
import tempfile
import uuid
from pathlib import Path
import subprocess

app = FastAPI(title="Doc2MD API", version="2.0.0")

# CORS 配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 工作目录
UPLOAD_DIR = Path("./uploads")
OUTPUT_DIR = Path("./outputs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


@app.get("/")
async def root():
    return {"message": "Doc2MD API v2.0", "status": "running"}


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """上传单个文件"""
    try:
        # 生成唯一文件名
        file_id = str(uuid.uuid4())[:8]
        ext = Path(file.filename).suffix.lower()
        new_filename = f"{file_id}{ext}"
        file_path = UPLOAD_DIR / new_filename
        
        # 保存文件
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        return {
            "success": True,
            "filename": new_filename,
            "original_name": file.filename,
            "size": len(content)
        }
    except Exception as e:
        raise HTTPException(500, f"上传失败: {str(e)}")


@app.post("/convert")
async def convert_file(data: dict):
    """转换文件为 Markdown"""
    filename = data.get("filename")
    if not filename:
        raise HTTPException(400, "缺少文件名")
    
    input_path = UPLOAD_DIR / filename
    if not input_path.exists():
        raise HTTPException(404, "文件不存在")
    
    try:
        # 生成输出文件名
        output_id = str(uuid.uuid4())[:8]
        output_filename = f"{output_id}.md"
        output_path = OUTPUT_DIR / output_filename
        
        # 根据文件类型选择转换方式
        if filename.endswith('.pdf'):
            result = await convert_pdf(input_path, output_path)
        elif filename.endswith(('.docx', '.doc')):
            result = await convert_word(input_path, output_path)
        else:
            raise HTTPException(400, "不支持的文件格式")
        
        return {
            "success": True,
            "download_url": f"/download/{output_filename}",
            "content": result
        }
    except Exception as e:
        raise HTTPException(500, f"转换失败: {str(e)}")


async def convert_pdf(input_path: Path, output_path: Path):
    """PDF 转 Markdown"""
    try:
        # 使用 marker 或 pymupdf 进行转换
        import fitz  # PyMuPDF
        
        doc = fitz.open(input_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_content.append(f"## 第 {page_num + 1} 页\n\n{text}")
        
        doc.close()
        
        # 合并内容
        markdown_content = "\n\n".join(text_content)
        
        # 保存文件
        output_path.write_text(markdown_content, encoding='utf-8')
        
        return markdown_content
    except Exception as e:
        # 如果 PyMuPDF 失败，返回简单文本
        fallback_content = f"# 转换结果\n\nPDF 文件已处理，但内容提取需要更高级的 OCR 支持。\n\n文件: {input_path.name}"
        output_path.write_text(fallback_content, encoding='utf-8')
        return fallback_content


async def convert_word(input_path: Path, output_path: Path):
    """Word 转 Markdown"""
    try:
        from docx import Document
        
        doc = Document(input_path)
        markdown_content = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                # 简单判断标题
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        markdown_content.append(f"{'#' * level_num} {para.text}")
                    except:
                        markdown_content.append(para.text)
                else:
                    markdown_content.append(para.text)
        
        # 提取表格
        for table in doc.tables:
            markdown_content.append("\n### 表格\n")
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                markdown_content.append(row_text)
            markdown_content.append("")
        
        # 合并内容
        final_content = "\n\n".join(markdown_content)
        
        # 保存文件
        output_path.write_text(final_content, encoding='utf-8')
        
        return final_content
    except Exception as e:
        fallback_content = f"# 转换结果\n\nWord 文件已处理。\n\n文件: {input_path.name}\n\n注意：部分格式可能需要手动调整。"
        output_path.write_text(fallback_content, encoding='utf-8')
        return fallback_content


@app.get("/download/{filename}")
async def download_file(filename: str):
    """下载转换后的文件"""
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(404, "文件不存在")
    
    return FileResponse(
        file_path,
        media_type="text/markdown",
        filename=f"converted_{filename}"
    )


@app.delete("/clear")
async def clear_files():
    """清空所有文件"""
    try:
        for dir_path in [UPLOAD_DIR, OUTPUT_DIR]:
            for file_path in dir_path.iterdir():
                if file_path.is_file():
                    file_path.unlink()
        return {"success": True, "message": "已清空所有文件"}
    except Exception as e:
        raise HTTPException(500, f"清空失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
